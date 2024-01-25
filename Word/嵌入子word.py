import os

import comtypes.client

tmp_path = r"C:\Users\12074\Desktop\测试\tmp_word-原始.docx"
# child_path = r"C:\Users\12074\Desktop\测试\full_csv_table0-副本.docx"
child_path = r"C:\Users\12074\Desktop\测试\full_csv_table0.docx"
child_path1 = r"C:\Users\12074\Desktop\测试\full_csv_table1.docx"
child_path2 = r"C:\Users\12074\Desktop\测试\full_csv_table2.docx"
save_path = r"C:\Users\12074\Desktop\测试\tmp_word-存储.docx"
from docx import Document
from docxcompose.composer import Composer


def insert_subdocuments(template_path, subdocuments, output_path):
    # 打开模板文档
    template_doc = Document(template_path)

    # 创建合成器
    composer = Composer(template_doc)

    # paragraph 偏移量
    subdocuments_offset = 0
    # 遍历每个子文档及其对应的占位符
    for placeholder, subdocument_path in subdocuments:
        # 打开子文档
        subdocument_doc = Document(subdocument_path)

        # 在文档中查找包含占位符的段落
        for i, paragraph in enumerate(template_doc.paragraphs):
            print(f"第{i}段落:{paragraph.text}")
            if placeholder in paragraph.text:
                print(f"匹配到字符串:{paragraph.text}")

                # paragraph.clear()
                # paragraph.text = ""
                # 清除段落的 runs 列表
                # for run in template_doc.paragraphs[i].runs:
                #     run.clear()
                # 直接删除段落对象
                del template_doc.paragraphs[i]

                composer.insert(i + subdocuments_offset, subdocument_doc)
                break
        # 保存合并后的文档
        composer.save(output_path)
        # template_doc = Document(output_path)
        # composer = Composer(template_doc)
        subdocuments_offset += 1


# 定义占位符
subdocuments = [
    ('{{full_csv_0}}', child_path),
    ('{{full_csv_1}}', child_path1),
    ('{{full_csv_2}}', child_path2)
]

# 将多个子文档插入到模板文档的相应占位符位置
insert_subdocuments(tmp_path, subdocuments, save_path)

# 生成PDf

pdf_file = os.path.join(r"C:\Users\12074\Desktop\测试", os.path.basename(save_path).replace('.docx', '.pdf'))
# 打开word应用程序
word = comtypes.client.CreateObject('Word.Application')
# 打开指定word文件
doc = word.Documents.Open(save_path)
# 保存并转换为pdf文档
doc.SaveAs(pdf_file, FileFormat=17)
# 关闭文档
doc.Close()
# 退出word
word.Quit()
