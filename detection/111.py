tmp_path = r"C:\Users\12074\Desktop\测试\tmp_word-原始.docx"
# child_path = r"C:\Users\12074\Desktop\测试\full_csv_table0-副本.docx"
child_path = r"C:\Users\12074\Desktop\测试\full_csv_table0.docx"
child_path1 = r"C:\Users\12074\Desktop\测试\full_csv_table1.docx"
save_path = r"C:\Users\12074\Desktop\测试\tmp_word-存储.docx"
from docx import Document
from docxcompose.composer import Composer

def find_placeholder_paragraphs(doc, placeholder):
    # 在文档中查找包含占位符的段落
    placeholder_paragraphs = []
    for i, paragraph in enumerate(doc.paragraphs):
        print(f"第{i}段落:{paragraph.text}")
        if placeholder in paragraph.text:
            paragraph.delete()
            print(f"匹配到字符串:{paragraph.text}")
            placeholder_paragraphs.append(i)
    return placeholder_paragraphs

def insert_subdocument(template_path, subdocument_path, output_path, placeholder):
    # 打开模板文档和子文档
    template_doc = Document(template_path)
    subdocument_doc = Document(subdocument_path)

    # 创建合成器
    composer = Composer(template_doc)

    # 查找包含占位符的段落
    placeholder_paragraphs = find_placeholder_paragraphs(template_doc, placeholder)

    if not placeholder_paragraphs:
        raise ValueError(f"Placeholder '{placeholder}' not found in the template.")

    # 在每个包含占位符的段落之前插入子文档的内容
    for paragraph_index in placeholder_paragraphs:
        # 获取占位符所在段落的前一个段落索引
        insertion_index = paragraph_index if paragraph_index > 0 else 0

        # 在合成器中插入子文档的内容
        composer.insert(insertion_index,subdocument_doc)

    # 保存合并后的文档
    composer.save(output_path)

def insert_subdocuments(template_path, subdocuments, output_path):
    # 打开模板文档
    template_doc = Document(template_path)

    # 创建合成器
    composer = Composer(template_doc)

    # 遍历每个子文档及其对应的占位符
    for subdocument_path, placeholder in subdocuments:
        # 打开子文档
        subdocument_doc = Document(subdocument_path)

        # 查找包含占位符的段落
        placeholder_paragraphs = find_placeholder_paragraphs(template_doc, placeholder)

        if not placeholder_paragraphs:
            raise ValueError(f"Placeholder '{placeholder}' not found in the template.")

        # 在每个包含占位符的段落之前插入子文档的内容
        for paragraph_index in placeholder_paragraphs:
            # 获取占位符所在段落的前一个段落索引
            insertion_index = paragraph_index if paragraph_index > 0 else 0

            # 在合成器中插入子文档的内容（根据 docxcompose 版本选择合适的写法）
            composer.insert(insertion_index, subdocument_doc)
        print("一个站位行结束\n")

    # 保存合并后的文档
    composer.save(output_path)

# 定义占位符
subdocuments = [
    (child_path, '{{full_csv_0}}'),
    (child_path1, '{{full_csv_1}}')
]


# 将多个子文档插入到模板文档的相应占位符位置
insert_subdocuments(tmp_path, subdocuments, save_path)
